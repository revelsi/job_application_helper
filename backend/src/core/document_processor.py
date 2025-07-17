"""
Copyright 2024 Job Application Helper Contributors

Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at

    http://www.apache.org/licenses/LICENSE-2.0

Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

"""
Simple document text extraction for job application helper.

Focuses on extracting text from PDFs and DOCX files using PyMuPDF as the primary extractor.
Designed to be simple, reliable, and easily extensible.
"""

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional, Union

try:
    import fitz  # PyMuPDF
    import pymupdf4llm

    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from src.utils.logging import get_logger

logger = get_logger(__name__)

# Constants
BYTES_PER_MB = 1024 * 1024
DEFAULT_MAX_FILE_SIZE_MB = 50


@dataclass
class DocumentContent:
    """Extracted content from a document."""

    text: str
    markdown: Optional[str] = None
    word_count: int = 0
    file_path: Optional[Path] = None
    extraction_time: Optional[float] = None
    success: bool = True
    error: Optional[str] = None

    def __post_init__(self):
        """Calculate word count if not provided."""
        if self.text and self.word_count == 0:
            self.word_count = len(self.text.split())


class DocumentProcessor:
    """Simple document text extractor focusing on PDFs and DOCX files."""

    @staticmethod
    def sanitize_personal_info(text: str) -> str:
        """Redact emails and phone numbers from text with improved precision."""
        if not text:
            return text
            
        original_length = len(text)
        
        # More precise email regex - requires proper email format
        # This avoids matching things like "file.txt" or "version.2"
        text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL]', text)
        
        # More precise phone number regex - matches common formats
        # US/International formats: (123) 456-7890, 123-456-7890, 123.456.7890, +1 123 456 7890
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b',  # US format
            r'\b\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b',  # Standard format
            r'\+[0-9]{1,3}[-.\s]?[0-9]{1,4}[-.\s]?[0-9]{1,4}[-.\s]?[0-9]{1,9}\b',  # International
        ]
        
        for pattern in phone_patterns:
            text = re.sub(pattern, '[PHONE]', text)
        
        # Log redaction impact
        final_length = len(text)
        if final_length < original_length:
            redacted_chars = original_length - final_length
            logger.debug(f"Personal info redaction: removed {redacted_chars} characters ({redacted_chars/original_length*100:.1f}%)")
        
        return text

    def __init__(self):
        """Initialize the processor and check available extractors."""
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which extractors are available."""
        available = ["Built-in (TXT)"]  # TXT is always available

        if HAS_PYMUPDF:
            available.append("PyMuPDF (PDF)")
        if HAS_DOCX:
            available.append("python-docx (DOCX)")

        if len(available) == 1:  # Only TXT available
            logger.warning(
                "Only TXT files supported. Install dependencies for more formats: pip install pymupdf pymupdf4llm python-docx"
            )

        logger.info(f"Document processor initialized with: {', '.join(available)}")

    def can_process(self, file_path: Path) -> bool:
        """Check if we can process this file type."""
        if not file_path.exists() or not file_path.is_file():
            return False

        suffix = file_path.suffix.lower()

        if suffix == ".pdf" and HAS_PYMUPDF:
            return True
        if suffix == ".docx" and HAS_DOCX:
            return True
        if suffix == ".txt":
            return True

        return False

    def extract_text(self, file_path: Union[str, Path]) -> DocumentContent:
        """
        Extract text from a document.

        Args:
            file_path: Path to the document file

        Returns:
            DocumentContent with extracted text and metadata
        """
        file_path = Path(file_path)
        start_time = datetime.now()

        # Validate file
        if not file_path.exists():
            return DocumentContent(
                text="",
                file_path=file_path,
                success=False,
                error=f"File not found: {file_path}",
            )

        if not self.can_process(file_path):
            return DocumentContent(
                text="",
                file_path=file_path,
                success=False,
                error=f"Unsupported file type: {file_path.suffix}",
            )

        # Extract based on file type
        try:
            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                content = self._extract_pdf(file_path)
            elif suffix == ".docx":
                content = self._extract_docx(file_path)
            elif suffix == ".txt":
                content = self._extract_txt(file_path)
            else:
                return DocumentContent(
                    text="",
                    file_path=file_path,
                    success=False,
                    error=f"Unsupported file type: {suffix}",
                )

            # Sanitize personal information in extracted text and markdown
            content.text = self.sanitize_personal_info(content.text)
            if content.markdown:
                content.markdown = self.sanitize_personal_info(content.markdown)

            # Add processing metadata
            content.file_path = file_path
            content.extraction_time = (datetime.now() - start_time).total_seconds()

            logger.info(f"Extracted {content.word_count} words from {file_path.name}")
            return content

        except Exception as e:
            logger.error(f"Failed to extract from {file_path}: {e}")
            return DocumentContent(
                text="", file_path=file_path, success=False, error=str(e)
            )

    def _extract_pdf(self, file_path: Path) -> DocumentContent:
        """Extract text from PDF using PyMuPDF with enhanced text cleaning."""
        if not HAS_PYMUPDF:
            raise RuntimeError("PyMuPDF not available")

        # Extract plain text
        doc = fitz.open(str(file_path))
        text_parts = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()

            # Clean up the page text
            page_text = self._clean_pdf_text(page_text)

            if page_text.strip():
                text_parts.append(page_text)

        plain_text = "\n\n".join(text_parts)
        doc.close()

        # Extract markdown (optimized for LLMs)
        markdown_text = None
        try:
            markdown_text = pymupdf4llm.to_markdown(str(file_path))
            # Clean the markdown text too
            if markdown_text:
                markdown_text = self._clean_pdf_text(markdown_text)
        except Exception as e:
            logger.warning(f"Markdown extraction failed: {e}")

        return DocumentContent(text=plain_text, markdown=markdown_text)

    def _clean_pdf_text(self, text: str) -> str:
        """Clean up PDF text to handle common extraction issues."""
        if not text:
            return text

        # Simple fix for common spaced text patterns
        # Only fix obvious cases where single capital letters are spaced
        text = re.sub(
            r"\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b", r"\1\2\3\4\5", text
        )
        text = re.sub(r"\b([A-Z])\s+([A-Z])\s+([A-Z])\s+([A-Z])\b", r"\1\2\3\4", text)
        text = re.sub(r"\b([A-Z])\s+([A-Z])\s+([A-Z])\b", r"\1\2\3", text)

        # General cleanup
        text = self._clean_extracted_text(text)

        return text

    def _extract_docx(self, file_path: Path) -> DocumentContent:
        """Extract text from DOCX using python-docx with enhanced table and textbox support."""
        if not HAS_DOCX:
            raise RuntimeError("python-docx not available")

        doc = DocxDocument(str(file_path))
        text_parts = []

        # Extract paragraphs (traditional text)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract tables (many modern resumes use tables for layout)
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text.strip():
                text_parts.append(table_text)

        # Extract headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                header_text = self._extract_header_footer_text(section.header)
                if header_text.strip():
                    text_parts.insert(0, header_text)  # Headers go at the top

            # Footer
            if section.footer:
                footer_text = self._extract_header_footer_text(section.footer)
                if footer_text.strip():
                    text_parts.append(footer_text)  # Footers go at the bottom

        # Extract text from textboxes and shapes (advanced layouts)
        # This is critical for modern resume templates
        try:
            textbox_text = self._extract_textboxes_and_shapes(doc)
            if textbox_text.strip():
                text_parts.append(textbox_text)
        except Exception as e:
            logger.warning(f"Failed to extract textboxes/shapes: {e}")

        # If we got very little text from normal extraction, try XML-based extraction
        combined_text = "\n\n".join(text_parts)
        if len(combined_text.strip()) < 100:  # Very little text extracted
            try:
                xml_text = self._extract_all_text_from_xml(doc)
                if xml_text.strip() and len(xml_text) > len(combined_text):
                    logger.info(
                        "Using XML-based extraction due to limited normal extraction"
                    )
                    combined_text = xml_text
            except Exception as e:
                logger.warning(f"XML-based extraction failed: {e}")

        # Clean up the text
        plain_text = self._clean_extracted_text(combined_text)

        return DocumentContent(text=plain_text)

    def _extract_textboxes_and_shapes(self, doc) -> str:
        """Extract text from textboxes and shapes using document XML."""
        try:
            from defusedxml import ElementTree as ET

            # Get the document XML
            doc_xml = doc._body._element.xml
            root = ET.fromstring(doc_xml)

            # Define namespaces for Word documents
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
                "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
                "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
                "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
                "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
            }

            textbox_texts = []

            # Look for text in various drawing elements
            # Search for text in textboxes (wps:txbx)
            for textbox in root.findall(".//wps:txbx", namespaces):
                for para in textbox.findall(".//w:p", namespaces):
                    para_text = self._extract_paragraph_text_from_xml(para, namespaces)
                    if para_text.strip():
                        textbox_texts.append(para_text.strip())

            # Search for text in shapes (a:t)
            for text_elem in root.findall(".//a:t", namespaces):
                if text_elem.text:
                    textbox_texts.append(text_elem.text.strip())

            # Search for text in drawing text elements (w:t within drawing)
            for drawing in root.findall(".//w:drawing", namespaces):
                for text_elem in drawing.findall(".//w:t", namespaces):
                    if text_elem.text:
                        textbox_texts.append(text_elem.text.strip())

            # Search for text in VML shapes (for older documents)
            for shape in root.findall(
                ".//v:shape", {"v": "urn:schemas-microsoft-com:vml"}
            ):
                for text_elem in shape.findall(".//w:t", namespaces):
                    if text_elem.text:
                        textbox_texts.append(text_elem.text.strip())

            return "\n".join(textbox_texts)

        except Exception as e:
            logger.debug(f"Advanced textbox extraction failed: {e}")
            return ""

    def _extract_paragraph_text_from_xml(self, para_elem, namespaces):
        """Extract text from a paragraph XML element."""
        text_parts = []
        for text_elem in para_elem.findall(".//w:t", namespaces):
            if text_elem.text:
                text_parts.append(text_elem.text)
        return "".join(text_parts)

    def _extract_table_text(self, table) -> str:
        """Extract text from a table with intelligent formatting and enhanced cell processing."""
        table_parts = []

        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()

                # If cell appears empty, try to extract from paragraphs directly
                if not cell_text:
                    para_texts = []
                    for para in cell.paragraphs:
                        para_text = para.text.strip()
                        if para_text:
                            para_texts.append(para_text)
                    cell_text = " ".join(para_texts)

                # If still empty, try to extract from runs
                if not cell_text:
                    run_texts = []
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text.strip()
                            if run_text:
                                run_texts.append(run_text)
                    cell_text = " ".join(run_texts)

                if cell_text:
                    # Handle nested tables in cells
                    if cell.tables:
                        for nested_table in cell.tables:
                            nested_text = self._extract_table_text(nested_table)
                            if nested_text.strip():
                                cell_text += "\n" + nested_text
                    row_cells.append(cell_text)

            if row_cells:
                # If it looks like a single-column layout, join with spaces
                # If it looks like multi-column data, use pipe separator
                if len(row_cells) == 1:
                    table_parts.append(row_cells[0])
                elif len(row_cells) <= 3 and all(len(cell) < 50 for cell in row_cells):
                    # Short cells, likely tabular data
                    table_parts.append(" | ".join(row_cells))
                else:
                    # Long cells, likely layout columns
                    table_parts.append(" ".join(row_cells))

        return "\n".join(table_parts)

    def _extract_header_footer_text(self, header_footer) -> str:
        """Extract text from headers and footers."""
        text_parts = []

        for paragraph in header_footer.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in header_footer.tables:
            table_text = self._extract_table_text(table)
            if table_text.strip():
                text_parts.append(table_text)

        return "\n".join(text_parts)

    def _clean_extracted_text(self, text: str) -> str:
        """Clean up extracted text to improve readability."""
        if not text:
            return text

        # Remove excessive whitespace
        import re

        # Replace multiple spaces with single space
        text = re.sub(r" +", " ", text)

        # Replace multiple newlines with double newlines
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)

        # Remove leading/trailing whitespace from each line
        lines = text.split("\n")
        cleaned_lines = [line.strip() for line in lines]

        # Remove empty lines at the beginning and end
        while cleaned_lines and not cleaned_lines[0]:
            cleaned_lines.pop(0)
        while cleaned_lines and not cleaned_lines[-1]:
            cleaned_lines.pop()

        return "\n".join(cleaned_lines)

    def _extract_txt(self, file_path: Path) -> DocumentContent:
        """Extract text from TXT file using UTF-8 encoding."""
        try:
            # Try UTF-8 first
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    text = f.read()
            except Exception as e:
                raise RuntimeError(
                    f"Failed to read text file with UTF-8 or latin-1 encoding: {e}"
                )

        return DocumentContent(text=text)

    def _extract_all_text_from_xml(self, doc) -> str:
        """Extract all text from the document XML as a fallback method."""
        try:
            from defusedxml import ElementTree as ET

            # Get the document XML
            doc_xml = doc._body._element.xml
            root = ET.fromstring(doc_xml)

            # Define namespace
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }

            # Extract all text elements
            text_elements = []
            for text_elem in root.findall(".//w:t", namespaces):
                if text_elem.text and text_elem.text.strip():
                    text_elements.append(text_elem.text.strip())

            # Join with spaces and clean up
            raw_text = " ".join(text_elements)

            # Try to add some structure by looking for patterns
            # This is a heuristic approach for better formatting
            structured_text = self._add_structure_to_xml_text(raw_text)

            return structured_text

        except Exception as e:
            logger.debug(f"XML text extraction failed: {e}")
            return ""

    def _add_structure_to_xml_text(self, text: str) -> str:
        """Add basic structure to XML-extracted text."""
        if not text:
            return text

        import re

        # Split into potential sections based on common patterns
        # Look for common resume section headers
        section_patterns = [
            r"\b(CONTACT|EDUCATION|EXPERIENCE|SKILLS|SUMMARY|PROFILE|OBJECTIVE|PROJECTS|CERTIFICATIONS|AWARDS|REFERENCES)\b",
            r"\b(Contact|Education|Experience|Skills|Summary|Profile|Objective|Projects|Certifications|Awards|References)\b",
            r"\b(contact|education|experience|skills|summary|profile|objective|projects|certifications|awards|references)\b",
        ]

        # Add line breaks before section headers
        for pattern in section_patterns:
            text = re.sub(pattern, r"\n\n\1", text)

        # Add line breaks before years (likely job dates)
        text = re.sub(r"\b(19|20)\d{2}\b", r"\n\1", text)

        # Add line breaks before email addresses
        text = re.sub(r"\b[\w\.-]+@[\w\.-]+\.\w+\b", r"\n\g<0>", text)

        # Add line breaks before phone numbers
        text = re.sub(r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b", r"\n\g<0>", text)

        # Clean up multiple newlines
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)

        return text.strip()

    def validate_file(
        self, file_path: Path, max_size_mb: int = DEFAULT_MAX_FILE_SIZE_MB
    ) -> tuple[bool, Optional[str]]:
        """
        Validate if a file can be processed.

        Args:
            file_path: Path to the file
            max_size_mb: Maximum file size in MB

        Returns:
            Tuple of (is_valid, error_message)
        """
        if not file_path.exists():
            return False, "File does not exist"

        if not file_path.is_file():
            return False, "Path is not a file"

        # Check file size
        file_size_mb = file_path.stat().st_size / BYTES_PER_MB
        if file_size_mb > max_size_mb:
            return False, f"File too large: {file_size_mb:.1f}MB (max: {max_size_mb}MB)"

        if not self.can_process(file_path):
            return False, f"Unsupported file type: {file_path.suffix}"

        return True, None


# Convenience function for simple usage
def extract_text_from_file(file_path: Union[str, Path]) -> str:
    """
    Simple function to extract text from a file.

    Returns empty string if extraction fails.
    """
    processor = DocumentProcessor()
    result = processor.extract_text(file_path)
    return result.text if result.success else ""
